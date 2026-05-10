"""
MediaFlow — Copilot attachments (v3.5.0-alpha.51).

Servizio per upload e processing di documenti caricati dal drawer copilot.
L'utente può allegare file (PDF, DOCX, TXT, MD, immagini) al suo messaggio
per dare contesto all'AI: capitolati cliente, brief, mail, screenshot, ecc.

Pipeline:
1. POST /ai/api/upload riceve il file → save_attachment salva su disk
2. extract_content estrae testo per i tipi text-based (pypdf/docx/raw)
   oppure ritorna metadata immagine
3. Il client riceve {file_id, filename, kind, extracted_text, ...}
4. Quando il client invia il prossimo messaggio, embed_attachments_in_text
   aggiunge gli extracted_text inline nel content user (formato leggibile
   per l'AI: header con nome file + contenuto troncato)

Storage: `uploads/copilot/{uuid}.{ext}` — file con nome originale preservato
nel manifest in-memory (non persistito in DB per MVP). Cleanup periodico
via cleanup_old_attachments() in lifespan.

NB: NON persistiamo in DB l'attachment per il MVP. Il file resta su disk,
il content estratto resta in memoria del client (incluso nei messaggi).
Per histories durature serve future estensione con AIAttachment model.
"""
from __future__ import annotations

import logging
import re
import uuid
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# Storage root — auto-created on first save
ATTACHMENT_DIR = Path("uploads") / "copilot"
ATTACHMENT_DIR.mkdir(parents=True, exist_ok=True)

# Limiti
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB
MAX_TEXT_CHARS = 50_000           # tronca text estratto a 50k chars (≈ ~12k tokens)
RETENTION_DAYS = 7

# Estensioni supportate
TEXT_EXTS = {".pdf", ".docx", ".txt", ".md"}
IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".gif"}
SUPPORTED_EXTS = TEXT_EXTS | IMAGE_EXTS

# v3.5.0-alpha.66.14.4 — Magic-bytes per validazione MIME oltre l'estensione.
# `txt` e `md` sono testo plain: nessuna firma, lasciamo passare se l'estensione
# matcha. Per gli altri tipi controlliamo i primi byte. WEBP richiede check su
# offset 8-12 in aggiunta al RIFF header (gestito custom).
MAGIC_BYTES: dict[str, list[bytes]] = {
    ".pdf": [b"%PDF-"],
    ".docx": [b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08"],  # ZIP container
    ".png": [b"\x89PNG\r\n\x1a\n"],
    ".jpg": [b"\xff\xd8\xff"],
    ".jpeg": [b"\xff\xd8\xff"],
    ".gif": [b"GIF87a", b"GIF89a"],
    # WEBP: gestito a parte, richiede 'RIFF????WEBP'
}


def _validate_magic_bytes(ext: str, content: bytes) -> None:
    """Solleva ValueError se i primi byte di `content` non corrispondono
    all'estensione dichiarata. Per .webp richiede check su byte 8-12."""
    if ext in (".txt", ".md"):
        return  # plain text: niente magic
    if ext == ".webp":
        if not (len(content) >= 12 and content[:4] == b"RIFF" and content[8:12] == b"WEBP"):
            raise ValueError(f"Contenuto non valido: l'header non corrisponde a {ext}")
        return
    expected = MAGIC_BYTES.get(ext)
    if not expected:
        return  # estensione gestita ma senza firma definita
    if not any(content.startswith(sig) for sig in expected):
        raise ValueError(f"Contenuto non valido: l'header non corrisponde a {ext}")


def save_attachment(filename: str, content: bytes, user_id: int) -> dict:
    """Salva un attachment caricato dall'utente. Ritorna metadata con file_id,
    kind ("text" o "image"), extracted_text (se text), preview info.

    v3.5.0-alpha.66.14.4 — `user_id` obbligatorio. Il file_id include il
    prefisso utente per ownership (`{user_id}-{uuid}`) — verificabile lato
    server senza dover persistere un manifest in DB.

    Solleva ValueError se: file troppo grande, estensione non supportata,
    magic-bytes incoerenti con l'estensione.
    """
    if not user_id or not isinstance(user_id, int):
        raise ValueError("user_id richiesto per upload attachment")
    if len(content) > MAX_FILE_SIZE:
        raise ValueError(f"File troppo grande ({len(content)} bytes, max {MAX_FILE_SIZE})")

    # Sanitize filename + estrai estensione
    safe_name = re.sub(r"[^\w\.\-]", "_", filename)[:120]
    ext = Path(safe_name).suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise ValueError(
            f"Estensione '{ext}' non supportata. Ammesse: "
            + ", ".join(sorted(SUPPORTED_EXTS))
        )

    # Magic bytes BEFORE write a disk: rifiutiamo prima di toccare il FS
    _validate_magic_bytes(ext, content)

    file_id = f"{user_id}-{uuid.uuid4().hex}"
    target = ATTACHMENT_DIR / f"{file_id}{ext}"
    target.write_bytes(content)

    kind = "image" if ext in IMAGE_EXTS else "text"

    extracted_text: Optional[str] = None
    extra: dict = {}

    if kind == "text":
        try:
            extracted_text = _extract_text(target, ext)
            if extracted_text and len(extracted_text) > MAX_TEXT_CHARS:
                extracted_text = extracted_text[:MAX_TEXT_CHARS] + (
                    f"\n\n[…troncato a {MAX_TEXT_CHARS} caratteri su {len(extracted_text)} totali]"
                )
        except Exception as e:
            logger.warning(f"Estrazione testo fallita per {filename}: {e}")
            extracted_text = f"[Errore estrazione testo: {e}]"
    else:
        # Immagine: leggi dimensioni se possibile
        try:
            from PIL import Image
            with Image.open(target) as img:
                extra["width"] = img.width
                extra["height"] = img.height
                extra["mode"] = img.mode
        except Exception as e:
            logger.warning(f"Lettura metadata immagine fallita: {e}")

    return {
        "file_id": file_id,
        "filename": safe_name,
        "ext": ext,
        "size": len(content),
        "kind": kind,
        "extracted_text": extracted_text,
        "extracted_text_chars": len(extracted_text) if extracted_text else 0,
        "url": f"/uploads/copilot/{file_id}{ext}",
        **extra,
    }


def _extract_text(path: Path, ext: str) -> str:
    """Estrae testo da PDF (pypdf), DOCX (python-docx), TXT/MD (raw)."""
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for i, p in enumerate(reader.pages):
            try:
                t = p.extract_text() or ""
                if t.strip():
                    pages.append(f"--- Pagina {i+1} ---\n{t.strip()}")
            except Exception as e:
                pages.append(f"--- Pagina {i+1} (errore): {e} ---")
        return "\n\n".join(pages)
    elif ext == ".docx":
        from docx import Document
        doc = Document(str(path))
        # Paragrafi + tabelle
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    elif ext in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="replace")
    else:
        return ""


def embed_attachments_in_text(user_text: str, attachments: list[dict]) -> str:
    """Costruisce il messaggio finale con gli allegati inline.

    Formato per attachment text:
      📎 ALLEGATO: filename.pdf (12345 chars)
      <contenuto estratto>
      ---FINE ALLEGATO---

    Per immagini (MVP α.51, no vision integration): nota placeholder.
    """
    if not attachments:
        return user_text
    parts = []
    for a in attachments:
        if a.get("kind") == "text" and a.get("extracted_text"):
            parts.append(
                f"📎 ALLEGATO: {a.get('filename', '?')} ({a.get('extracted_text_chars', 0)} caratteri)\n"
                f"{a['extracted_text']}\n"
                f"---FINE ALLEGATO---"
            )
        elif a.get("kind") == "image":
            # MVP: l'immagine è salvata ma non passata al provider (richiederebbe
            # vision blocks per Anthropic/OpenAI/Gemini). Da estendere in α.52.
            dims = ""
            if a.get("width") and a.get("height"):
                dims = f" {a['width']}x{a['height']}"
            parts.append(
                f"📎 ALLEGATO IMMAGINE: {a.get('filename', '?')}{dims}\n"
                f"[L'utente ha allegato un'immagine. URL: {a.get('url', '?')}. "
                f"Vision integration in arrivo nelle prossime versioni. "
                f"Chiedi all'utente di descrivere il contenuto se necessario.]\n"
                f"---FINE ALLEGATO---"
            )
        elif a.get("kind") == "text" and not a.get("extracted_text"):
            parts.append(
                f"📎 ALLEGATO: {a.get('filename', '?')} [vuoto o non leggibile]\n"
                f"---FINE ALLEGATO---"
            )
    if parts:
        return "\n\n".join(parts) + "\n\n" + user_text
    return user_text


def _media_type_for_ext(ext: str) -> str:
    """Mappa estensione → MIME type per Anthropic image blocks."""
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
    }.get(ext.lower(), "image/png")


def _ownership_ok(file_id: str, user_id: int) -> bool:
    """v3.5.0-alpha.66.14.4 — Verifica che `file_id` appartenga a `user_id`.

    Convenzione: `{user_id}-{uuid32}` (vedi save_attachment). Se il prefisso
    non matcha, l'utente sta tentando di linkare un file altrui.
    """
    if not file_id or not user_id:
        return False
    expected_prefix = f"{user_id}-"
    return file_id.startswith(expected_prefix)


def _make_image_block_from_attachment(att: dict, user_id: Optional[int] = None) -> Optional[dict]:
    """Costruisce un blocco image canonico Anthropic per un attachment.

    v3.5.0-alpha.53 — Vision integration. Ritorna None se il file non
    esiste più o non è un'immagine.
    v3.5.0-alpha.66.14.4 — Se `user_id` è passato, valida ownership: ritorna
    None se il file_id non appartiene all'utente corrente (anti-leak).
    Output formato Anthropic Messages API:
        {"type": "image", "source": {"type": "base64",
         "media_type": "image/png", "data": "..."}}
    """
    import base64
    if att.get("kind") != "image":
        return None
    file_id = att.get("file_id")
    ext = att.get("ext")
    if not file_id or not ext:
        return None
    if user_id is not None and not _ownership_ok(file_id, user_id):
        logger.warning(f"image attachment ownership denied: file_id={file_id} user_id={user_id}")
        return None
    p = ATTACHMENT_DIR / f"{file_id}{ext}"
    if not p.exists():
        return None
    try:
        data = p.read_bytes()
    except Exception as e:
        logger.warning(f"image read failed for {p}: {e}")
        return None
    # Anthropic ha cap a ~5MB per immagine, 20MB totali per richiesta.
    # Se >5MB ridimensiona/skippa
    if len(data) > 5 * 1024 * 1024:
        logger.warning(f"image {file_id}{ext} > 5MB, skipping")
        return None
    return {
        "type": "image",
        "source": {
            "type": "base64",
            "media_type": _media_type_for_ext(ext),
            "data": base64.b64encode(data).decode("ascii"),
        },
    }


def build_user_content_blocks(
    user_text: str,
    attachments: list[dict],
    supports_vision: bool,
    user_id: Optional[int] = None,
) -> "str | list[dict]":
    """Costruisce il content del messaggio user.

    v3.5.0-alpha.53 — Vision integration:
    - Se nessun allegato → ritorna `user_text` (stringa).
    - Se solo allegati testuali → embed inline (stringa, retrocompat).
    - Se allegati immagine + provider supports_vision → ritorna content
      list canonico Anthropic con text block (testo + extract testo) +
      image blocks (base64). Anthropic SDK lo accetta nativamente; gli
      altri provider lo traducono nel proprio formato (vedi
      OpenAIProvider).
    - Se allegati immagine + provider NON supports_vision → fallback
      placeholder testuale (comportamento α.51).

    NB: Se l'immagine è > 5MB o file scomparso, cade in placeholder
    testuale per quell'immagine specifica (gli altri restano image blocks).
    """
    if not attachments:
        return user_text

    # Separazione per tipo
    text_atts = [a for a in attachments if a.get("kind") == "text"]
    image_atts = [a for a in attachments if a.get("kind") == "image"]

    # Se non ci sono immagini OR il provider non supporta vision → fallback
    # alla vecchia funzione embed (placeholder per le immagini).
    if not image_atts or not supports_vision:
        return embed_attachments_in_text(user_text, attachments)

    # Path vision: text block consolidato + image blocks
    text_parts: list[str] = []
    for a in text_atts:
        if a.get("extracted_text"):
            text_parts.append(
                f"📎 ALLEGATO: {a.get('filename', '?')} "
                f"({a.get('extracted_text_chars', 0)} caratteri)\n"
                f"{a['extracted_text']}\n"
                f"---FINE ALLEGATO---"
            )
        else:
            text_parts.append(
                f"📎 ALLEGATO: {a.get('filename', '?')} [vuoto o non leggibile]\n"
                f"---FINE ALLEGATO---"
            )

    image_blocks: list[dict] = []
    fallback_text_for_failed_images: list[str] = []
    for a in image_atts:
        block = _make_image_block_from_attachment(a, user_id=user_id)
        if block is not None:
            # Aggiungi un testo prima dell'immagine per identificarla nel contesto
            text_parts.append(f"📎 IMMAGINE: {a.get('filename', '?')}")
            image_blocks.append(block)
        else:
            dims = ""
            if a.get("width") and a.get("height"):
                dims = f" {a['width']}x{a['height']}"
            fallback_text_for_failed_images.append(
                f"📎 IMMAGINE: {a.get('filename', '?')}{dims} [non caricabile]"
            )
    text_parts.extend(fallback_text_for_failed_images)

    if user_text:
        text_parts.append(user_text)

    blocks: list[dict] = []
    if text_parts:
        blocks.append({"type": "text", "text": "\n\n".join(text_parts)})
    blocks.extend(image_blocks)
    return blocks


def cleanup_old_attachments() -> int:
    """Elimina file più vecchi di RETENTION_DAYS giorni. Idempotente.
    Chiamato dal lifespan periodico dell'app. Ritorna conteggio eliminati."""
    cutoff = datetime.utcnow() - timedelta(days=RETENTION_DAYS)
    deleted = 0
    if not ATTACHMENT_DIR.exists():
        return 0
    for f in ATTACHMENT_DIR.iterdir():
        if not f.is_file():
            continue
        try:
            mtime = datetime.utcfromtimestamp(f.stat().st_mtime)
            if mtime < cutoff:
                f.unlink()
                deleted += 1
        except Exception as e:
            logger.warning(f"cleanup_old_attachments: skip {f.name}: {e}")
    if deleted:
        logger.info(f"cleanup_old_attachments: eliminati {deleted} file > {RETENTION_DAYS}gg")
    return deleted
