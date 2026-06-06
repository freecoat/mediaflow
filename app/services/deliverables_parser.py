"""
MediaFlow — Deliverables Parser
Estrae voci di capitolato da PDF / Word / Excel / testo libero
e le matcha con il listino prezzi tramite AI.
"""
from __future__ import annotations
import io, json, logging
from pathlib import Path
from typing import Optional
from app.services.ai_provider import get_provider
from app.services.naming_resolver import normalize_naming_convention

logger = logging.getLogger(__name__)


# ── Estrazione testo dai vari formati ───────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Estrai anche contenuto tabelle
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
        rows_text = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text.append(f"=== Foglio: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    rows_text.append(" | ".join(cells))
        return "\n".join(rows_text)
    except Exception as e:
        logger.error(f"XLSX extraction failed: {e}")
        return ""


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    elif ext in (".xlsx", ".xls"):
        return extract_text_from_xlsx(file_bytes)
    elif ext in (".txt", ".md"):
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""
    return ""


# ── Prompt per parsing capitolato ───────────────────────────

PARSE_SYSTEM_PROMPT = """Sei un assistente esperto in postproduzione audiovisiva (cinema, TV, pubblicità).

Il tuo compito: analizzare un capitolato di consegne di un cliente e tradurlo in un elenco strutturato di voci operative.

Per ogni consegna/lavorazione richiesta, estrai:
- description: descrizione tecnica breve (es. "DCP 4K Mastering VF")
- detail: eventuale dettaglio aggiuntivo (es. "incl. QC e subtitling")
- quantity: quantità numerica (es. 1, 2, 90)
- unit: unità di misura (day/hr/min/pc/week/TB/GB/m/allow)
- category: categoria di postproduzione (PICTURE/SOUND/VFX/MASTERING/DELIVERABLES DCI/DELIVERABLES SOUND/DAILIES/ARCHIVE/TRANSFER/MATERIALS)
- section: sezione suggerita per raggruppamento in preventivo (A=Mastering/Delivery, B=Picture/Grading, C=Sound, D=VFX, E=Archive/Transfer)
- confidence: tuo livello di confidenza (high/medium/low)
- notes: eventuali ambiguità o info mancanti che l'utente deve verificare

Se il capitolato menziona durata del prodotto (es. "film 90 minuti"), usala per calcolare quantità per le voci "per min" (es. H264 encoding).

Se il capitolato è vago, fai ipotesi ragionevoli ma indica confidence:"low" e spiega in notes.

Schema output:
{
  "project_info": {
    "title": "...",
    "client": "...",
    "length_minutes": 90,
    "fps": "24",
    "delivery_format": "4K DCI",
    "shooting_format": "ARRI Alexa",
    "delivery_deadline": "2024-12-31"
  },
  "deliverables": [
    {
      "description": "4K DCP Mastering VF",
      "detail": "incl. full QC - EN Version",
      "quantity": 1,
      "unit": "pc",
      "category": "MASTERING",
      "section": "A",
      "confidence": "high",
      "notes": null
    },
    ...
  ],
  "global_notes": "Note generali sul capitolato nel suo insieme"
}"""


def parse_deliverables(text: str, hint: Optional[str] = None, provider=None) -> Optional[dict]:
    """
    Analizza un capitolato e restituisce la struttura voci + info progetto.
    hint: testo addizionale dell'utente (es. "è un documentario da 52 minuti").
    provider: instance opzionale (per-utente). Se None usa get_provider() global.
    v3.5.0-alpha.172.81 (Bundle F): accetta provider iniettato dal router
    per usare la AI key per-utente invece del fallback global.
    """
    if provider is None:
        provider = get_provider()
    if not provider:
        logger.warning("AI provider non disponibile — parser disabilitato")
        return None

    if len(text.strip()) < 20:
        logger.warning("Testo troppo breve per il parsing")
        return None

    # Limita lunghezza testo per non saturare il context
    MAX_CHARS = 30000
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[... testo troncato ...]"

    user_prompt = f"""Capitolato da analizzare:

---
{text}
---
"""
    if hint:
        user_prompt += f"\nInfo aggiuntiva dall'utente: {hint}\n"

    result = provider.extract_json(PARSE_SYSTEM_PROMPT, user_prompt, max_tokens=6000)
    return result


# ── Estrazione blocchi DeliveryTemplate (v3.5.0-alpha.66.20 F14) ─

PARSE_TEMPLATE_SYSTEM_PROMPT = """Sei un esperto di postproduzione e capitolati di consegna (delivery schedule) per cinema/TV/streaming.

Compito: leggere un capitolato e produrre **8 blocchi JSON strutturati** che descrivono le specifiche tecniche del template di consegna.

I 8 blocchi (tutti opzionali, ometti i campi mancanti):

1. video_specs: { codec, container, resolution, fps, scan, colorspace, bit_depth, gamma, white_point, hdr, prores_flavor, dpx_pad, ... }
2. audio_specs: { config, channel_layout, sample_rate, bit_depth, peak_dbfs, lufs_target, true_peak, dialnorm, languages, codec, ... }
3. text_specs: { subtitles, closed_captions, forced_narratives, languages, format (SCC/SRT/IMSC/STL), embed_or_sidecar, ... }
4. head_format: { bars_color, bars_duration_s, slate_required, slate_layout, beep_2pop, head_silence_s, timecode_start, ... }
5. textless_format: { required, type (clean/dirty), reels, formats, ... }
6. naming_convention (oggetto, opzionale — compila SOLO se il capitolato specifica una convenzione di nomenclatura file; altrimenti ometti o metti null):
   - "pattern": stringa con token tra graffe scelti TRA QUESTI: {project_code, project_title, film_name, content_type, aspect, resolution, framerate, audio_config, lang_audio, lang_subs, territory, version, revision, standard, package_type, deliverable_kind, date_iso, date_compact, studio_code, facility_code}. Esempio: "{film_name}_{content_type}_{resolution}_{lang_audio}_{date_compact}".
   - "separator": separatore (es. "_").
   - "case": "upper" | "lower" | "asis".
   - "extension": estensione file se indicata (es. ".mxf", ".wav") o "".
   - "max_length": numero massimo caratteri se indicato, altrimenti null.
   - "allowed_chars": classe caratteri ammessi se indicata (es. "A-Za-z0-9_-").
   - "examples": lista di nomi-file di esempio citati nel capitolato.
   - "raw_note": se la convenzione è descritta a parole ma NON mappabile a un pattern pulito, riporta qui il testo verbatim.
   Estrai questo blocco SIA per il capitolato nel suo insieme (chiave "naming_convention" a livello root) SIA, quando il capitolato distingue per singola consegna, dentro ogni voce dell'array "deliverables" (campo "naming_convention" omogeneo allo stesso schema).
7. archive_specs: { master_format, lto_generation, hash_required, redundancy, media_type, ... }
8. metadata_requirements: { mxf_metadata, exif_required, xml_sidecar (IMF/IMSC), iso639, fps_metadata, dolbyvision_xml, ... }

Inoltre:
- name: nome sintetico (es. "Netflix HDR10 IMF v1.3")
- broadcaster: emittente/distributore (es. "Netflix", "A24", "Sky")
- code: codice breve uppercase (es. "NETFLIX-IMF-HDR10")
- description: 1-2 frasi
- ai_confidence: 0..1 (quanto sei sicuro dell'estrazione)

Schema output:
{
  "code": "NETFLIX-IMF",
  "name": "Netflix IMF Standard 1.3",
  "broadcaster": "Netflix",
  "description": "Master IMF per delivery Netflix originals (HDR10/Atmos).",
  "video_specs": {...},
  "audio_specs": {...},
  "text_specs": {...},
  "head_format": {...},
  "textless_format": {...},
  "naming_convention": {"pattern": "{film_name}_{content_type}_{resolution}_{lang_audio}_{date_compact}", "separator": "_", "case": "upper", "extension": ".mxf", "max_length": 120, "allowed_chars": "A-Za-z0-9_-", "examples": [...], "raw_note": ""},
  "archive_specs": {...},
  "metadata_requirements": {...},
  "deliverables": [
    {"name": "DCP 4K VF", "naming_convention": {"pattern": "{film_name}_FTR_{resolution}_{lang_audio}", "separator": "_", "case": "upper"}}
  ],
  "ai_confidence": 0.85
}

Il blocco "deliverables" è opzionale: includilo solo quando il capitolato distingue una convenzione di naming per singola consegna (ogni voce porta il proprio "naming_convention" allo stesso schema del blocco 6). Se un blocco non è menzionato nel capitolato, ometti la chiave. Non inventare specifiche assenti."""


def parse_delivery_template(text: str, provider=None) -> Optional[dict]:
    """Analizza un capitolato e ritorna un dict con i 8 blocchi DeliveryTemplate
    + metadati (code/name/broadcaster/description/ai_confidence).

    Usato dall'endpoint POST /delivery-templates/api/parse per popolare la
    preview prima del salvataggio. L'utente può poi correggere/integrare
    prima di salvare. v3.5.0-alpha.66.20 Fase 2 step C.
    v3.5.0-alpha.172.81 (Bundle F): accetta provider iniettato dal router.
    """
    if provider is None:
        provider = get_provider()
    if not provider:
        logger.warning("AI provider non disponibile — parse_delivery_template disabilitato")
        return None
    if len(text.strip()) < 20:
        return None
    MAX_CHARS = 30000
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[... testo troncato ...]"
    user_prompt = f"""Capitolato da analizzare:

---
{text}
---

Estrai i blocchi strutturati come da schema."""
    # v3.5.0-alpha.172.111 — max_tokens 4000→8000 per capitolati grossi.
    # Senza questo bump A24/IRDA tornavano JSON troncato a metà struct,
    # safe_json_parse falliva → 503 "Risposta non JSON valido".
    result = provider.extract_json(PARSE_TEMPLATE_SYSTEM_PROMPT, user_prompt, max_tokens=8000)
    if not isinstance(result, dict):
        return result
    # α.172.182 (NC-T4) — normalizza la naming convention grezza dell'AI prima
    # del save (template + eventuali override per-item). normalize_* ritorna None
    # se l'AI ha omesso/lasciato vuoto il blocco → safe assegnare comunque.
    result["naming_convention"] = normalize_naming_convention(result.get("naming_convention"))
    items = result.get("deliverables")
    if isinstance(items, list):
        for it in items:
            if isinstance(it, dict):
                it["naming_convention"] = normalize_naming_convention(it.get("naming_convention"))
    # Catena capitolato→fisico: deriva requires_physical/physical_media_kind
    # dall'archive_specs a livello di template e li inietta sia nel risultato
    # root (per chi salva il template) sia in ogni voce deliverables (usata da
    # materialize_items in delivery_items_parser.py).
    _rp, _pmk = derive_physical_from_archive_specs(result.get("archive_specs"))
    result["requires_physical"] = _rp
    result["physical_media_kind"] = _pmk
    if isinstance(items, list):
        for it in items:
            if isinstance(it, dict):
                # Eredita dal template; se l'item avesse specs proprie future
                # si potrebbe overridare qui — per ora propagazione diretta.
                it.setdefault("requires_physical", _rp)
                it.setdefault("physical_media_kind", _pmk)
    return result


# ── Heuristica archivio fisico (catena capitolato→fisico) ───

_PHYSICAL_MEDIA_ALLOWED = {"lto", "hdd", "cru", "bluray", "dvd", "case", "other"}


def derive_physical_from_archive_specs(archive_specs: Optional[dict]) -> tuple[bool, Optional[str]]:
    """Deriva (requires_physical, physical_media_kind) dall'archive_specs del template.

    Strategia difensiva: tutte le chiavi opzionali, valori possono essere str/list/dict.
    Costruisce un blob testuale da tutti i valori scalari (anche annidati) e applica
    matching per keyword.

    Valori di ritorno:
      (True,  "lto")    se menziona LTO/LTFS
      (True,  "cru")    se menziona CRU
      (True,  "hdd")    se menziona HDD/hard drive/hard disk/drive
      (True,  "bluray") se menziona Blu-ray/Bluray/BD25/BD50
      (True,  "dvd")    se menziona DVD
      (False, None)     se nessun medium fisico riconosciuto
    """
    if not archive_specs or not isinstance(archive_specs, dict):
        return False, None

    def _collect_text(obj) -> str:
        """Appiattisce ricorsivamente un dict/list in stringa lowercase."""
        if isinstance(obj, str):
            return obj.lower()
        if isinstance(obj, (int, float)):
            return str(obj).lower()
        if isinstance(obj, list):
            return " ".join(_collect_text(v) for v in obj)
        if isinstance(obj, dict):
            return " ".join(_collect_text(v) for v in obj.values())
        return ""

    blob = _collect_text(archive_specs)

    if "lto" in blob or "ltfs" in blob:
        return True, "lto"
    if "cru" in blob:
        return True, "cru"
    if "hdd" in blob or "hard drive" in blob or "hard disk" in blob:
        return True, "hdd"
    # "drive" da solo è troppo generico (es. "drive link"), controlliamo solo
    # se accompagnato da altri indizi già esclusi sopra — lasciamo fuori.
    if "blu-ray" in blob or "bluray" in blob or "bd25" in blob or "bd50" in blob:
        return True, "bluray"
    if "dvd" in blob:
        return True, "dvd"

    return False, None


# ── Matching voci capitolato ↔ listino prezzi ───────────────

MATCH_SYSTEM_PROMPT = """Sei un assistente che mappa voci di capitolato di un cliente sul listino prezzi interno della casa di postproduzione.

Dato un elenco di voci da capitolato e un listino prezzi, per ogni voce del capitolato:
1. Trova la voce di listino più simile (se esiste)
2. Restituisci l'id del match migliore, oppure null se nessuna voce è abbinabile
3. Indica il livello di confidenza del match

Criteri di matching:
- Terminologia tecnica (DCP, DCDM, ProRes, H264, grading, mix, conform, etc.)
- Unità di misura compatibile (day/min/pc)
- Categoria coerente (SOUND→SOUND, PICTURE→PICTURE)

Schema output:
{
  "matches": [
    {
      "deliverable_index": 0,
      "price_item_id": 42,
      "confidence": "high",
      "reasoning": "Match esatto: '4K DCP Mastering' = listino #42 '4K DCP Mastering VF'"
    },
    ...
  ]
}

Se una voce non ha corrispondenza chiara nel listino, imposta price_item_id: null e spiega in reasoning."""


def match_deliverables_to_pricelist(deliverables: list[dict],
                                     pricelist: list[dict],
                                     provider=None) -> Optional[dict]:
    """
    Matcha le voci di capitolato con le voci del listino.
    deliverables: output di parse_deliverables['deliverables']
    pricelist: lista voci dal DB con campi id, name, category, unit, price_list
    v3.5.0-alpha.172.81 (Bundle F): accetta provider iniettato dal router.
    """
    if provider is None:
        provider = get_provider()
    if not provider:
        return None

    deliverables_text = json.dumps(
        [{"index": i, **d} for i, d in enumerate(deliverables)],
        ensure_ascii=False, indent=2)
    pricelist_text = json.dumps(
        [{"id": p["id"], "name": p["name"], "category": p.get("category"),
          "unit": p["unit"], "price_list": p.get("price_list")} for p in pricelist],
        ensure_ascii=False, indent=2)

    user_prompt = f"""Voci capitolato da matchare:
{deliverables_text}

Listino prezzi disponibile:
{pricelist_text}

Genera le mappature."""

    result = provider.extract_json(MATCH_SYSTEM_PROMPT, user_prompt, max_tokens=4000)
    return result
